import json
import shutil
import tempfile
import unittest
from datetime import datetime, timedelta
from pathlib import Path

from docx import Document

from platform_app.main import (
    build_project_context,
    build_cache_artifact_paths,
    build_execution_status_items,
    build_upload_storage_name,
    build_project_file_lookup,
    enrich_basis_items,
    get_milestone_status_label,
    get_runtime_system_status,
    serialize_project_messages,
    run_deepseek_connectivity_test,
    sanitize_review_update_data,
    serialize_pending_review_entries,
    serialize_project_file_entries,
    sync_project_requirements_from_payload,
)
from platform_app.models import ExtractionField, ExtractionJob, Project, ProjectFile, ProjectMessage, ProjectMilestone, ProjectRequirement, SystemSetting
from platform_app.services import ai_pipeline


class FakeQuery:
    def __init__(self, rows):
        self.rows = rows

    def filter(self, *args, **kwargs):
        return self

    def all(self):
        return list(self.rows)


class FakeSession:
    def __init__(self, existing_rows):
        self.existing_rows = list(existing_rows)
        self.added = []
        self.settings = {}

    def query(self, model):
        if model is ProjectRequirement:
            return FakeQuery(self.existing_rows)
        raise AssertionError(f"Unexpected model: {model}")

    def add(self, item):
        self.added.append(item)

    def get(self, model, key):
        if model is SystemSetting:
            return self.settings.get(key)
        raise AssertionError(f"Unexpected model: {model}")


class ProjectFileFlowTests(unittest.TestCase):
    def test_build_upload_storage_name_preserves_name_and_uniqueness_shape(self) -> None:
        first = build_upload_storage_name("same-name.docx")
        second = build_upload_storage_name("same-name.docx")

        self.assertTrue(first.endswith("-same-name.docx"))
        self.assertTrue(second.endswith("-same-name.docx"))
        self.assertNotEqual(first, second)

    def test_build_cache_artifact_paths_uses_unique_stem(self) -> None:
        temp_dir = Path(tempfile.mkdtemp())
        try:
            first = build_cache_artifact_paths(temp_dir, "20260709120000123-same-name.docx")
            second = build_cache_artifact_paths(temp_dir, "20260709120000999-same-name.docx")

            self.assertNotEqual(first["text_path"], second["text_path"])
            self.assertTrue(first["text_path"].name.endswith(".txt"))
            self.assertTrue(second["parse_path"].name.endswith(".json"))
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

    def test_docx_extraction_preserves_paragraph_table_order(self) -> None:
        temp_dir = Path(tempfile.mkdtemp())
        try:
            docx_path = temp_dir / "order-test.docx"
            document = Document()
            document.add_paragraph("项目名称：云庭物业服务项目")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "开标地点"
            table.cell(0, 1).text = "温州市公共资源交易中心"
            document.add_paragraph("递交注意事项：正本一份，副本四份。")
            document.save(docx_path)

            extracted = ai_pipeline.extract_document_text(docx_path, docx_path.name)

            self.assertLess(extracted.index("项目名称"), extracted.index("开标地点"))
            self.assertLess(extracted.index("开标地点"), extracted.index("递交注意事项"))
            self.assertIn("[表格 1 | 第 1 行]", extracted)
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

    def test_ai_fields_merge_values_back_into_project_payload(self) -> None:
        extracted_text = "\n".join(
            [
                "[段落 1] 项目名称：云庭物业服务项目",
                "[段落 2] 递交注意事项：投标文件正本一份，副本四份，递交时携带授权委托书和身份证。",
                "[段落 3] 封标要求：封套封口处加盖公章，封皮注明项目名称和投标人名称。",
                "[段落 4] 废标条款：未按要求密封的投标文件作无效投标处理。",
                "[段落 5] 本项目需要项目负责人现场答辩。",
            ]
        )
        response_payload = {
            "summary": "测试解析",
            "project": {"name": "云庭物业服务项目"},
            "fields": [
                {
                    "field_key": "submission_notes",
                    "field_label": "递交注意事项",
                    "extracted_value": "投标文件正本一份，副本四份，递交时携带授权委托书和身份证。",
                    "source_location": "段落 2",
                    "source_quote": "递交注意事项：投标文件正本一份，副本四份，递交时携带授权委托书和身份证。",
                    "confidence": "high",
                },
                {
                    "field_key": "seal_notes",
                    "field_label": "封标注意事项",
                    "extracted_value": "封套封口处加盖公章，封皮注明项目名称和投标人名称。",
                    "source_location": "段落 3",
                    "source_quote": "封标要求：封套封口处加盖公章，封皮注明项目名称和投标人名称。",
                    "confidence": "high",
                },
            ],
            "requirements": [],
        }

        original = ai_pipeline.deepseek_chat
        ai_pipeline.deepseek_chat = lambda *args, **kwargs: json.dumps(response_payload, ensure_ascii=False)
        try:
            project, fields, _requirements, _summary = ai_pipeline.ai_extract_fields("云庭.docx", extracted_text, api_key="sk-test")
        finally:
            ai_pipeline.deepseek_chat = original

        field_map = {item["field_key"]: item for item in fields}
        self.assertIn("正本一份", project["submission_notes"])
        self.assertIn("封套封口", project["seal_notes"])
        self.assertIn("授权委托书", field_map["submission_notes"]["source_quote"])
        self.assertEqual(field_map["submission_notes"]["confidence"], "high")

    def test_ai_extract_fields_falls_back_when_deepseek_returns_bad_json(self) -> None:
        extracted_text = "\n".join(
            [
                "[段落 1] 项目名称：云庭物业服务项目",
                "[段落 2] 投标保证金：5万",
                "[段落 3] 递交注意事项：正本一份，副本四份。",
                "[段落 4] 废标条款：未按要求密封的投标文件作无效投标处理。",
            ]
        )

        original = ai_pipeline.deepseek_chat
        ai_pipeline.deepseek_chat = lambda *args, **kwargs: '{"summary":"broken","project":{"name":"云庭物业服务项目"'
        try:
            project, fields, requirements, summary = ai_pipeline.ai_extract_fields("云庭.docx", extracted_text, api_key="sk-test")
        finally:
            ai_pipeline.deepseek_chat = original

        self.assertIn("云庭物业服务项目", project["name"])
        self.assertIn("DeepSeek 返回 JSON 格式异常", summary)
        self.assertTrue(any(item["field_key"] == "invalidation_risks" for item in fields))
        self.assertTrue(any(item["category"] == "risk" for item in requirements))

    def test_sync_project_requirements_adds_only_new_grounded_items(self) -> None:
        existing = ProjectRequirement(
            project_id=1,
            category="qualification",
            title="营业执照",
            content="需提供有效营业执照复印件",
            importance="high",
            source_location="第 3 页",
            source_quote="需提供有效营业执照复印件",
            created_by="tester",
        )
        session = FakeSession([existing])

        added_count, skipped_count, duplicate_count = sync_project_requirements_from_payload(
            session,
            project_id=1,
            requirement_payloads=[
                {
                    "category": "qualification",
                    "title": "营业执照",
                    "content": "需提供有效营业执照复印件",
                    "importance": "high",
                    "source_location": "第 3 页",
                    "source_quote": "需提供有效营业执照复印件",
                },
                {
                    "category": "document",
                    "title": "投标文件组成",
                    "content": "商务标、技术标、报价文件分别装订。",
                    "importance": "medium",
                    "source_location": "第 8 页",
                    "source_quote": "投标文件应包括商务标、技术标、报价文件。",
                },
                {
                    "category": "risk",
                    "title": "废标条款",
                    "content": "未按要求盖章视为无效投标。",
                    "importance": "high",
                    "source_location": "",
                    "source_quote": "",
                },
            ],
            created_by="tester",
        )

        self.assertEqual(added_count, 1)
        self.assertEqual(skipped_count, 1)
        self.assertEqual(duplicate_count, 1)
        self.assertEqual(len(session.added), 1)
        self.assertEqual(session.added[0].title, "投标文件组成")

    def test_sanitize_review_update_data_keeps_existing_master_fields(self) -> None:
        project = Project(
            name="正式项目名称",
            short_name="正式项目",
            tender_code="FORMAL-001",
            buyer="正式业主",
            owner_name="正式负责人",
            bid_mode="self",
            status="tracking",
        )
        approved_data = {
            "name": "补充文件里的名称",
            "short_name": "补充简称",
            "tender_code": "SUP-001",
            "buyer": "补充业主",
            "owner_name": "补充负责人",
            "location": "杭州市西湖区",
            "service_scope": "补充服务范围",
        }
        fields = [
            ExtractionField(field_key="name", status="confirmed", final_value="补充文件里的名称"),
            ExtractionField(field_key="short_name", status="confirmed", final_value="补充简称"),
            ExtractionField(field_key="tender_code", status="confirmed", final_value="SUP-001"),
            ExtractionField(field_key="buyer", status="confirmed", final_value="补充业主"),
            ExtractionField(field_key="owner_name", status="confirmed", final_value="补充负责人"),
            ExtractionField(field_key="location", status="confirmed", final_value="杭州市西湖区"),
        ]

        sanitized = sanitize_review_update_data(project, approved_data, fields)

        self.assertNotIn("name", sanitized)
        self.assertNotIn("short_name", sanitized)
        self.assertNotIn("tender_code", sanitized)
        self.assertNotIn("buyer", sanitized)
        self.assertNotIn("owner_name", sanitized)
        self.assertEqual(sanitized["location"], "杭州市西湖区")
        self.assertEqual(sanitized["service_scope"], "补充服务范围")

    def test_serialize_project_file_entries_marks_intake_and_supplement(self) -> None:
        now = datetime.utcnow()
        intake_file = ProjectFile(
            id=1,
            project_id=9,
            original_name="main-tender.docx",
            storage_path="storage/projects/9/source/main-tender.docx",
            file_hash="hash-1",
            file_size=2048,
            created_by="Alice",
            created_at=now,
        )
        supplement_file = ProjectFile(
            id=2,
            project_id=9,
            original_name="clarification.pdf",
            storage_path="storage/projects/9/source/clarification.pdf",
            file_hash="hash-2",
            file_size=4096,
            source_kind="supplement",
            created_by="Bob",
            created_at=now + timedelta(minutes=10),
        )

        entries = serialize_project_file_entries([supplement_file, intake_file], 9)

        by_name = {item["original_name"]: item for item in entries}
        self.assertEqual(by_name["main-tender.docx"]["kind"], "intake")
        self.assertEqual(by_name["clarification.pdf"]["kind"], "supplement")
        self.assertEqual(by_name["clarification.pdf"]["state"], "active")
        self.assertEqual(by_name["clarification.pdf"]["download_url"], "/projects/9/files/2")

    def test_enrich_basis_items_links_to_project_file_and_kind(self) -> None:
        file_record = ProjectFile(
            id=7,
            project_id=3,
            original_name="sample_tender.txt",
            storage_path="storage/projects/3/source/sample_tender.txt",
            file_hash="hash-7",
            file_size=1024,
            created_at=datetime.utcnow(),
        )
        lookup = build_project_file_lookup([file_record])

        enriched = enrich_basis_items(
            [
                {
                    "file_name": "sample_tender.txt",
                    "source_location": "第 6 行",
                    "source_quote": "投标保证金：5万",
                    "explanation": "以下原文直接支持本次回答。",
                },
                {
                    "file_name": "",
                    "source_location": "历史来源说明",
                    "source_quote": "旧版本引用信息",
                    "explanation": "这条问答创建于结构化依据上线前。",
                },
            ],
            lookup,
            3,
        )

        self.assertEqual(enriched[0]["download_url"], "/projects/3/files/7")
        self.assertEqual(enriched[0]["basis_kind"], "quote")
        self.assertEqual(enriched[1]["basis_kind"], "history")
        self.assertEqual(enriched[1]["download_url"], "")

    def test_serialize_pending_review_entries_marks_supplement(self) -> None:
        file_record = ProjectFile(
            id=5,
            project_id=None,
            original_name="qa-appendix.pdf",
            storage_path="storage/projects/pending/project-11/source/qa-appendix.pdf",
            file_hash="hash-5",
            file_size=4096,
            source_kind="supplement",
            created_by="Carol",
            created_at=datetime.utcnow(),
        )
        job = ExtractionJob(
            id=15,
            project_id=None,
            matched_project_id=11,
            project_file_id=5,
            project_file=file_record,
            created_by="Carol",
            confidence_summary="提取完成，待确认",
            created_at=datetime.utcnow(),
        )

        entries = serialize_pending_review_entries([job], 11)

        self.assertEqual(entries[0]["kind"], "supplement")
        self.assertEqual(entries[0]["review_url"], "/reviews/15")
        self.assertEqual(entries[0]["ai_mode_label"], "未标记")

    def test_serialize_pending_review_entries_exposes_ai_mode_label(self) -> None:
        file_record = ProjectFile(
            id=6,
            project_id=None,
            original_name="intake.docx",
            storage_path="storage/projects/pending/202607/source/intake.docx",
            file_hash="hash-6",
            file_size=2048,
            created_by="Dora",
            created_at=datetime.utcnow(),
        )
        job = ExtractionJob(
            id=16,
            project_id=None,
            matched_project_id=None,
            project_file_id=6,
            project_file=file_record,
            created_by="Dora",
            confidence_summary="当前未配置 DeepSeek，已使用规则兜底提取。",
            ai_model="规则兜底",
            created_at=datetime.utcnow(),
        )

        entries = serialize_pending_review_entries([job], 99)

        self.assertEqual(entries[0]["ai_mode_label"], "规则兜底")

    def test_get_runtime_system_status_reads_database_backed_flags(self) -> None:
        session = FakeSession([])
        session.settings["deepseek_api_key"] = SystemSetting(key="deepseek_api_key", value="sk-test")
        session.settings["paddleocr_api_url"] = SystemSetting(key="paddleocr_api_url", value="https://ocr.example/api")
        session.settings["paddleocr_access_token"] = SystemSetting(key="paddleocr_access_token", value="token-123")
        session.settings["paddleocr_timeout_seconds"] = SystemSetting(key="paddleocr_timeout_seconds", value="300")

        status = get_runtime_system_status(session)

        self.assertTrue(status["ai_configured"])
        self.assertEqual(status["ai_source"], "database")
        self.assertTrue(status["ocr_configured"])
        self.assertEqual(status["ocr_source"], "database")
        self.assertTrue(status["scan_pdf_enabled"])

    def test_get_runtime_system_status_marks_missing_when_no_ai_key(self) -> None:
        session = FakeSession([])

        status = get_runtime_system_status(session)

        self.assertFalse(status["ai_configured"])
        self.assertEqual(status["ai_source"], "missing")

    def test_run_deepseek_connectivity_test_uses_database_key(self) -> None:
        session = FakeSession([])
        session.settings["deepseek_api_key"] = SystemSetting(key="deepseek_api_key", value="sk-db-value")

        import platform_app.main as main_module

        original = main_module.ping_deepseek
        captured: list[str | None] = []

        def fake_ping(api_key: str | None = None) -> dict[str, str]:
            captured.append(api_key)
            return {"status": "ok", "model": "deepseek-v4-flash", "base_url": "https://api.deepseek.com"}

        main_module.ping_deepseek = fake_ping
        try:
            result = run_deepseek_connectivity_test(session)
        finally:
            main_module.ping_deepseek = original

        self.assertEqual(captured, ["sk-db-value"])
        self.assertEqual(result["status"], "ok")

    def test_serialize_project_messages_exposes_answer_mode(self) -> None:
        message = ProjectMessage(
            id=9,
            project_id=3,
            question="这个项目的保证金是什么？",
            answer="投标保证金：5万",
            citations='{"version":2,"answer_mode":"fallback","answer_status":"grounded","basis_items":[],"ai_suggestion":""}',
            created_by="Alice",
            created_at=datetime.utcnow(),
        )

        items = serialize_project_messages([message], file_lookup={}, project_id=3)

        self.assertEqual(items[0]["answer_mode"], "fallback")
        self.assertEqual(items[0]["answer_mode_label"], "规则兜底")

    def test_get_milestone_status_label_uses_execution_specific_copy(self) -> None:
        self.assertEqual(get_milestone_status_label("deposit", "pending"), "未汇出")
        self.assertEqual(get_milestone_status_label("deposit", "done"), "已汇出")
        self.assertEqual(get_milestone_status_label("custom", "pending"), "未开始")

    def test_build_execution_status_items_orders_and_labels_milestones(self) -> None:
        project = Project(name="测试项目", bid_mode="self", status="tracking")
        now = datetime.utcnow()
        items = build_execution_status_items(
            project,
            [
                ProjectMilestone(
                    id=3,
                    project_id=1,
                    milestone_type="submission",
                    title="递交投标文件",
                    due_at=now,
                    status="pending",
                    created_at=now,
                    updated_at=now,
                ),
                ProjectMilestone(
                    id=2,
                    project_id=1,
                    milestone_type="deposit",
                    title="汇出保证金",
                    due_at=now,
                    status="done",
                    created_at=now,
                    updated_at=now,
                ),
            ],
        )

        self.assertEqual([item["milestone_type"] for item in items], ["deposit", "submission"])
        self.assertEqual(items[0]["status_label"], "已汇出")
        self.assertEqual(items[1]["status_label"], "未递交")


if __name__ == "__main__":
    unittest.main()
