from __future__ import annotations

import json
import os
import re
import urllib.error
import urllib.request
from datetime import datetime
from pathlib import Path
from typing import Any

import requests

try:  # Optional legacy-only document stack; the Hermes workflow does not load these routes.
    import pdfplumber
    from docx import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from pypdf import PdfReader
except ModuleNotFoundError:  # pragma: no cover - exercised by core-only deployment smoke tests.
    pdfplumber = None
    Document = None
    CT_Tbl = CT_P = Table = Paragraph = None
    PdfReader = None

from ..config import (
    get_deepseek_base_url,
    get_deepseek_model,
    get_paddleocr_access_token_env_key,
    get_paddleocr_api_url_env_key,
    get_paddleocr_timeout_env_key,
    get_system_setting_env_key,
)


MAX_MODEL_CONTEXT_CHARS = 120_000
MODEL_CONTEXT_HEAD_CHARS = 30_000
MODEL_CONTEXT_TAIL_CHARS = 12_000

FIELD_LABELS = {
    "name": "项目名称",
    "short_name": "项目简称",
    "tender_code": "招标编号",
    "buyer": "招标人/采购人",
    "project_type": "项目类型",
    "bid_mode": "投标性质",
    "owner_name": "负责人",
    "agency": "代理机构",
    "contact_name": "联系人",
    "contact_phone": "联系电话",
    "location": "地点",
    "service_scope": "服务范围",
    "contract_term": "合同期限",
    "budget_amount": "预算金额/控制价",
    "deposit_amount": "保证金金额",
    "signup_deadline": "报名截止",
    "document_sale_deadline": "文件购买截止",
    "clarification_deadline": "疑问澄清截止",
    "site_visit_time": "踏勘时间",
    "deposit_deadline": "保证金截止",
    "bid_datetime": "开标时间",
    "submission_datetime": "递交时间",
    "bid_location": "开标/递交地点",
    "file_fee": "文件费",
    "payment_info": "收款信息",
    "has_defense": "是否有答辩",
    "defense_presenter": "答辩人选",
    "invalidation_risks": "废标项/无效投标风险",
    "submission_notes": "递交注意事项",
    "seal_notes": "封标注意事项",
    "notes": "备注",
}

REQUIREMENT_CATEGORY_LABELS = {
    "qualification": "资格要求",
    "document": "文件组成",
    "business": "商务要求",
    "technical": "技术要求",
    "pricing": "报价要求",
    "scoring": "评分办法",
    "risk": "废标/风险",
    "response": "关键响应点",
    "other": "其他要求",
}

FIELD_SOURCE_KEYWORDS: dict[str, list[str]] = {
    "name": ["项目名称", "采购项目名称", "招标项目名称"],
    "short_name": ["项目简称"],
    "tender_code": ["项目编号", "招标编号", "采购编号", "招标文件编号"],
    "buyer": ["招标人", "采购人", "采购单位", "建设单位"],
    "project_type": ["项目类型", "采购方式", "招标方式"],
    "bid_mode": ["投标性质", "联合体", "是否接受联合体"],
    "owner_name": ["项目负责人", "经办人", "联系人"],
    "agency": ["代理机构", "采购代理机构", "招标代理"],
    "contact_name": ["联系人", "项目联系人"],
    "contact_phone": ["联系电话", "联系方式", "电话"],
    "location": ["项目地点", "服务地点", "履约地点", "地址"],
    "service_scope": ["服务范围", "采购内容", "服务内容", "项目概况"],
    "contract_term": ["服务期限", "合同期限", "履约期限", "服务期"],
    "budget_amount": ["预算金额", "最高限价", "控制价", "采购预算"],
    "deposit_amount": ["投标保证金", "保证金金额", "保证金"],
    "signup_deadline": ["报名截止", "报名时间", "获取招标文件"],
    "document_sale_deadline": ["文件购买截止", "获取招标文件截止", "发售截止"],
    "clarification_deadline": ["疑问澄清截止", "答疑截止", "提疑截止", "质疑截止"],
    "site_visit_time": ["踏勘时间", "现场踏勘", "踏勘"],
    "deposit_deadline": ["保证金截止", "保证金缴纳截止", "保证金到账", "投标保证金"],
    "bid_datetime": ["开标时间", "开标日期", "投标截止时间"],
    "submission_datetime": ["投标文件递交截止时间", "递交截止时间", "投标截止时间"],
    "bid_location": ["开标地点", "递交地点", "投标地点"],
    "file_fee": ["文件费", "招标文件售价", "售价"],
    "payment_info": ["收款信息", "账户信息", "汇款信息", "开户行", "账号"],
    "has_defense": ["答辩", "述标", "演示", "方案汇报"],
    "defense_presenter": ["答辩人", "述标人", "项目负责人答辩", "项目经理答辩"],
    "invalidation_risks": ["废标", "无效投标", "否决投标", "投标无效", "无效响应", "不予受理"],
    "submission_notes": ["递交", "投标文件递交", "送达", "正本", "副本", "份", "携带", "授权委托", "身份证"],
    "seal_notes": ["密封", "封套", "封条", "封皮", "封标", "加盖公章", "骑缝章"],
    "notes": ["其他要求", "特别说明", "注意事项"],
}

REQUIREMENT_SOURCE_KEYWORDS: dict[str, list[str]] = {
    "qualification": ["资格要求", "投标人资格", "申请人资格", "供应商资格"],
    "document": ["投标文件组成", "响应文件组成", "投标文件包括", "资格证明材料", "投标文件格式", "附件", "附表"],
    "business": ["商务要求", "商务条款", "合同条款", "付款方式"],
    "technical": ["技术要求", "服务要求", "服务方案", "技术标准"],
    "pricing": ["报价要求", "报价方式", "报价文件", "最高限价"],
    "scoring": ["评分办法", "评审办法", "评标办法", "综合评分"],
    "risk": ["废标", "无效投标", "否决投标", "投标无效", "无效响应"],
    "response": ["响应", "实质性要求", "必须满足", "不得偏离"],
    "other": ["注意事项", "特别说明", "其他要求"],
}

DOCUMENT_CONTEXT_KEYWORDS = sorted(
    {
        keyword
        for keywords in [*FIELD_SOURCE_KEYWORDS.values(), *REQUIREMENT_SOURCE_KEYWORDS.values()]
        for keyword in keywords
    }
    | {"附件", "附表", "投标文件格式", "响应文件格式", "密封", "封装", "递交", "开标", "答辩"}
)

QUESTION_ANSWER_STATUS_LABELS = {
    "grounded": "已按文件回答",
    "partial": "部分依据",
    "not_found": "文件未找到",
    "unavailable": "AI 未配置",
}

ANSWER_MODE_LABELS = {
    "deepseek": "DeepSeek",
    "fallback": "规则兜底",
    "unknown": "历史记录",
}

DEFAULT_PADDLEOCR_TIMEOUT = 600.0

LOCAL_QA_GROUPS = [
    {"label": "项目名称", "mode": "value", "hints": ["项目名称", "项目名"]},
    {"label": "招标编号", "mode": "value", "hints": ["招标编号", "项目编号", "采购编号", "编号"]},
    {"label": "招标人/采购人", "mode": "value", "hints": ["招标人", "采购人", "甲方"]},
    {"label": "项目地点", "mode": "value", "hints": ["项目地点", "服务地点", "地点", "区域"]},
    {"label": "报名截止时间", "mode": "value", "hints": ["报名截止", "报名时间"]},
    {"label": "投标保证金", "mode": "value", "hints": ["投标保证金", "保证金", "保证金金额"]},
    {"label": "保证金截止时间", "mode": "value", "hints": ["保证金截止", "保证金缴纳截止", "保证金时间"]},
    {"label": "投标文件递交截止时间", "mode": "value", "hints": ["递交截止", "投标截止", "递交时间", "投标文件递交"]},
    {"label": "开标时间", "mode": "value", "hints": ["开标时间", "开标日期", "投标时间"]},
    {"label": "疑问澄清截止时间", "mode": "value", "hints": ["疑问澄清", "答疑截止", "疑问截止"]},
    {"label": "文件购买截止时间", "mode": "value", "hints": ["文件购买", "发售截止", "获取招标文件"]},
    {"label": "踏勘时间", "mode": "value", "hints": ["踏勘", "现场踏勘"]},
    {"label": "投标文件组成", "mode": "snippet", "hints": ["投标文件组成", "文件组成", "商务标", "技术标", "资格证明文件"]},
    {"label": "资格要求", "mode": "snippet", "hints": ["资格要求", "资质要求", "业绩要求", "人员要求"]},
    {"label": "评分办法", "mode": "snippet", "hints": ["评分办法", "评标办法", "评分标准", "打分"]},
    {"label": "报价要求", "mode": "snippet", "hints": ["报价", "价格", "最高限价", "控制价"]},
    {"label": "废标条款", "mode": "snippet", "hints": ["废标", "无效投标", "否决投标", "风险条款"]},
]


def clean_text(value: object) -> str | None:
    if value is None:
        return None
    text = str(value).strip()
    return text or None


def count_cjk_chars(text: str) -> int:
    return sum(1 for char in text if "\u4e00" <= char <= "\u9fff")


def repair_mojibake_text(text: str) -> str:
    if not text:
        return text
    try:
        repaired = text.encode("latin1").decode("utf-8")
    except (UnicodeEncodeError, UnicodeDecodeError):
        return text
    if repaired != text and count_cjk_chars(repaired) > count_cjk_chars(text):
        return repaired
    return text


def get_runtime_ai_settings(api_key: str | None = None) -> dict[str, object]:
    resolved_key = api_key or os.getenv(get_system_setting_env_key()) or os.getenv("DEEPSEEK_API_KEY")
    return {
        "api_key": resolved_key,
        "base_url": get_deepseek_base_url(),
        "model": get_deepseek_model(),
        "configured": bool(clean_text(resolved_key)),
    }


def get_runtime_ai_mode_label(api_key: str | None = None) -> tuple[str, str]:
    settings = get_runtime_ai_settings(api_key)
    if settings["configured"]:
        return "deepseek", str(settings.get("model") or "DeepSeek")
    return "fallback", "规则兜底"


def get_runtime_ocr_settings(
    api_url: str | None = None,
    access_token: str | None = None,
    timeout_seconds: str | float | None = None,
) -> dict[str, object]:
    resolved_api_url = clean_text(api_url) or os.getenv(get_paddleocr_api_url_env_key()) or ""
    resolved_token = clean_text(access_token) or os.getenv(get_paddleocr_access_token_env_key()) or ""
    resolved_timeout = timeout_seconds
    if resolved_timeout is None:
        resolved_timeout = os.getenv(get_paddleocr_timeout_env_key())

    timeout_value = DEFAULT_PADDLEOCR_TIMEOUT
    if resolved_timeout not in (None, ""):
        try:
            parsed_timeout = float(resolved_timeout)
            if parsed_timeout > 0:
                timeout_value = parsed_timeout
        except (TypeError, ValueError):
            timeout_value = DEFAULT_PADDLEOCR_TIMEOUT

    return {
        "api_url": resolved_api_url.strip(),
        "access_token": resolved_token.strip(),
        "timeout_seconds": timeout_value,
        "configured": bool(resolved_api_url.strip() and resolved_token.strip()),
    }


def parse_ai_datetime(value: object) -> str | None:
    text = clean_text(value)
    if not text:
        return None

    text = text.replace("年", "-").replace("月", "-").replace("日", " ")
    text = re.sub(r"[：]", ":", text)
    text = re.sub(r"\s+", " ", text).strip()

    patterns = [
        r"(\d{4})[-/.](\d{1,2})[-/.](\d{1,2})\s+(\d{1,2}):(\d{1,2})",
        r"(\d{4})[-/.](\d{1,2})[-/.](\d{1,2})",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if not match:
            continue

        parts = [int(item) for item in match.groups()]
        if len(parts) == 3:
            year, month, day = parts
            hour, minute = 0, 0
        else:
            year, month, day, hour, minute = parts

        try:
            return datetime(year, month, day, hour, minute).strftime("%Y-%m-%d %H:%M")
        except ValueError:
            return None
    return None


def normalize_ai_project_payload(payload: dict[str, object]) -> dict[str, object]:
    def text_or_empty(name: str) -> str:
        return clean_text(payload.get(name)) or ""

    deposit_amount = payload.get("deposit_amount")
    if isinstance(deposit_amount, str):
        amount_match = re.search(r"\d+(?:\.\d+)?", deposit_amount)
        deposit_amount = amount_match.group(0) if amount_match else None

    bid_mode = clean_text(payload.get("bid_mode")) or "self"
    if bid_mode not in {"self", "partner"}:
        bid_mode = "self"

    has_defense = text_or_empty("has_defense")
    defense_aliases = {
        "需要": "yes",
        "有": "yes",
        "是": "yes",
        "yes": "yes",
        "true": "yes",
        "无需": "no",
        "不需要": "no",
        "无": "no",
        "否": "no",
        "no": "no",
        "false": "no",
    }
    normalized_has_defense = defense_aliases.get(has_defense.lower(), defense_aliases.get(has_defense, ""))

    return {
        "name": text_or_empty("name") or "未命名投标项目",
        "short_name": text_or_empty("short_name"),
        "tender_code": text_or_empty("tender_code"),
        "buyer": text_or_empty("buyer"),
        "project_type": text_or_empty("project_type"),
        "bid_mode": bid_mode,
        "owner_name": text_or_empty("owner_name") or "待分配",
        "agency": text_or_empty("agency"),
        "contact_name": text_or_empty("contact_name"),
        "contact_phone": text_or_empty("contact_phone"),
        "location": text_or_empty("location"),
        "service_scope": text_or_empty("service_scope"),
        "contract_term": text_or_empty("contract_term"),
        "budget_amount": text_or_empty("budget_amount"),
        "deposit_amount": str(deposit_amount).strip() if deposit_amount not in (None, "") else "",
        "signup_deadline": parse_ai_datetime(payload.get("signup_deadline")),
        "document_sale_deadline": parse_ai_datetime(payload.get("document_sale_deadline")),
        "clarification_deadline": parse_ai_datetime(payload.get("clarification_deadline")),
        "site_visit_time": parse_ai_datetime(payload.get("site_visit_time")),
        "deposit_deadline": parse_ai_datetime(payload.get("deposit_deadline")),
        "bid_datetime": parse_ai_datetime(payload.get("bid_datetime")),
        "submission_datetime": parse_ai_datetime(payload.get("submission_datetime")),
        "bid_location": text_or_empty("bid_location"),
        "file_fee": text_or_empty("file_fee"),
        "payment_info": text_or_empty("payment_info"),
        "has_defense": normalized_has_defense,
        "defense_presenter": text_or_empty("defense_presenter"),
        "invalidation_risks": text_or_empty("invalidation_risks"),
        "submission_notes": text_or_empty("submission_notes"),
        "seal_notes": text_or_empty("seal_notes"),
        "notes": text_or_empty("notes"),
    }


def quote_is_grounded_in_text(source_quote: str, text: str) -> bool:
    quote = clean_text(source_quote) or ""
    if not quote:
        return False
    if len(quote) <= 8:
        return True
    normalized_quote = re.sub(r"\s+", "", quote)
    normalized_text = re.sub(r"\s+", "", text)
    return normalized_quote[:160] in normalized_text


def merge_project_payload_from_fields(project_payload: object, fields: object) -> dict[str, object]:
    merged = dict(project_payload) if isinstance(project_payload, dict) else {}
    if not isinstance(fields, list):
        return merged

    for item in fields:
        if not isinstance(item, dict):
            continue
        field_key = clean_text(item.get("field_key"))
        if field_key not in FIELD_LABELS:
            continue
        if clean_text(merged.get(field_key)):
            continue
        extracted_value = clean_text(item.get("extracted_value"))
        if extracted_value:
            merged[field_key] = extracted_value
    return merged


def normalize_ai_field_payload(
    items: object,
    project_payload: dict[str, object],
    extracted_text: str,
) -> list[dict[str, str]]:
    by_key: dict[str, dict[str, str]] = {}
    if isinstance(items, list):
        for item in items:
            if not isinstance(item, dict):
                continue
            field_key = clean_text(item.get("field_key"))
            if field_key not in FIELD_LABELS:
                continue
            by_key[field_key] = {
                "field_key": field_key,
                "field_label": clean_text(item.get("field_label")) or FIELD_LABELS[field_key],
                "extracted_value": clean_text(item.get("extracted_value")) or "",
                "source_location": clean_text(item.get("source_location")) or "",
                "source_quote": clean_text(item.get("source_quote")) or "",
                "confidence": clean_text(item.get("confidence")) or "",
            }

    fallback_sources = build_fallback_field_sources(extracted_text, project_payload)
    normalized: list[dict[str, str]] = []
    for field_key, field_label in FIELD_LABELS.items():
        item = by_key.get(field_key, {})
        value = clean_text(item.get("extracted_value")) or clean_text(project_payload.get(field_key)) or ""
        source_location = clean_text(item.get("source_location")) or ""
        source_quote = clean_text(item.get("source_quote")) or ""

        if value and (not source_quote or not quote_is_grounded_in_text(source_quote, extracted_text)):
            source_meta = fallback_sources.get(field_key, {})
            source_location = clean_text(source_meta.get("source_location")) or source_location
            source_quote = clean_text(source_meta.get("source_quote")) or source_quote

        if value and source_quote and not quote_is_grounded_in_text(source_quote, extracted_text):
            source_location = ""
            source_quote = ""

        confidence = clean_text(item.get("confidence")) or ("medium" if source_quote else "low")
        normalized.append(
            {
                "field_key": field_key,
                "field_label": field_label,
                "extracted_value": value,
                "source_location": source_location,
                "source_quote": source_quote,
                "confidence": confidence,
            }
        )
    return normalized


def normalize_requirement_payload(items: object) -> list[dict[str, str]]:
    normalized: list[dict[str, str]] = []
    if not isinstance(items, list):
        return normalized

    for item in items:
        if not isinstance(item, dict):
            continue
        title = clean_text(item.get("title"))
        content = clean_text(item.get("content"))
        if not title and not content:
            continue
        category = clean_text(item.get("category")) or "other"
        if category not in REQUIREMENT_CATEGORY_LABELS:
            category = "other"
        importance = clean_text(item.get("importance")) or "medium"
        if importance not in {"high", "medium", "low"}:
            importance = "medium"
        normalized.append(
            {
                "category": category,
                "category_label": REQUIREMENT_CATEGORY_LABELS.get(category, category),
                "title": title or content or "未命名要求",
                "content": content or title or "",
                "importance": importance,
                "source_location": clean_text(item.get("source_location")) or "",
                "source_quote": clean_text(item.get("source_quote")) or "",
            }
        )
    return normalized


def normalize_answer_basis_items(items: object) -> list[dict[str, str]]:
    normalized: list[dict[str, str]] = []
    if not isinstance(items, list):
        return normalized

    for item in items:
        if not isinstance(item, dict):
            continue
        file_name = clean_text(item.get("file_name")) or ""
        source_location = clean_text(item.get("source_location")) or ""
        source_quote = clean_text(item.get("source_quote")) or ""
        explanation = clean_text(item.get("explanation")) or ""
        if not any((file_name, source_location, source_quote, explanation)):
            continue
        normalized.append(
            {
                "file_name": file_name,
                "source_location": source_location,
                "source_quote": source_quote,
                "explanation": explanation,
            }
        )
    return normalized[:6]


def normalize_question_answer_payload(payload: dict[str, object]) -> dict[str, object]:
    answer_status = clean_text(payload.get("answer_status")) or "grounded"
    if answer_status not in QUESTION_ANSWER_STATUS_LABELS:
        answer_status = "grounded"

    answer_mode = clean_text(payload.get("answer_mode")) or "unknown"
    if answer_mode not in ANSWER_MODE_LABELS:
        answer_mode = "unknown"

    answer = clean_text(payload.get("answer")) or ""
    ai_suggestion = clean_text(payload.get("ai_suggestion")) or ""
    basis_items = normalize_answer_basis_items(payload.get("basis_items"))

    if answer_status == "grounded" and not basis_items:
        answer_status = "partial"

    if answer_status == "not_found" and not answer:
        answer = "当前已上传文件中未找到明确依据。"
    if answer_status == "unavailable" and not answer:
        answer = "当前未配置 DeepSeek API Key，暂时无法进行项目问答。"
    if answer_status in {"grounded", "partial"} and not answer:
        answer = "当前已根据已上传文件整理可用结论。"

    if answer_status == "not_found" and not ai_suggestion:
        ai_suggestion = "可以补传答疑文件、附件或换一个更具体的问题后再试。"

    return {
        "version": 2,
        "answer_mode": answer_mode,
        "answer_status": answer_status,
        "answer": answer,
        "basis_items": basis_items,
        "ai_suggestion": ai_suggestion,
    }


def serialize_question_answer_metadata(payload: dict[str, object]) -> str:
    normalized = normalize_question_answer_payload(payload)
    metadata = {
        "version": normalized["version"],
        "answer_mode": normalized["answer_mode"],
        "answer_status": normalized["answer_status"],
        "basis_items": normalized["basis_items"],
        "ai_suggestion": normalized["ai_suggestion"],
    }
    return json.dumps(metadata, ensure_ascii=False)


def parse_question_answer_metadata(answer: str, citations: str | None) -> dict[str, object]:
    raw_citations = clean_text(citations) or ""
    if raw_citations:
        try:
            payload = json.loads(raw_citations)
        except json.JSONDecodeError:
            payload = None
        if isinstance(payload, dict) and payload.get("version") == 2:
            payload["answer"] = clean_text(answer) or str(payload.get("answer", "") or "")
            return normalize_question_answer_payload(payload)

    answer_text = clean_text(answer) or "当前没有可展示的回答。"
    if raw_citations and "依据：" in answer_text:
        answer_text = answer_text.split("依据：", 1)[0].strip() or answer_text

    basis_items: list[dict[str, str]] = []
    if raw_citations:
        basis_items.append(
            {
                "file_name": "",
                "source_location": "历史来源说明",
                "source_quote": raw_citations,
                "explanation": "这条问答创建于结构化依据上线前，以下为原始来源说明。",
            }
        )

    answer_status = "grounded" if basis_items else "partial"
    if "未找到" in answer_text or "无法" in answer_text:
        answer_status = "not_found"

    return normalize_question_answer_payload(
        {
            "answer_mode": "unknown",
            "answer_status": answer_status,
            "answer": answer_text,
            "basis_items": basis_items,
            "ai_suggestion": "",
        }
    )


def extract_ocr_text(path: Path, original_name: str, ocr_settings: dict[str, object] | None = None) -> str:
    settings = ocr_settings or get_runtime_ocr_settings()
    if not settings.get("configured"):
        raise ValueError(
            "当前 PDF 疑似为扫描件，且未配置 PaddleOCR 文档解析。"
            "请先在系统设置中填写 PaddleOCR 接口地址和 Access Token 后重试。"
        )

    api_url = str(settings.get("api_url") or "").strip()
    access_token = str(settings.get("access_token") or "").strip()
    timeout_seconds = float(settings.get("timeout_seconds") or DEFAULT_PADDLEOCR_TIMEOUT)

    suffix = Path(original_name).suffix.lower()
    file_type = 0 if suffix == ".pdf" else 1
    payload: dict[str, Any] = {
        "file": path.read_bytes().hex(),
    }
    try:
        import base64

        payload["file"] = base64.b64encode(path.read_bytes()).decode("utf-8")
    except Exception as exc:
        raise ValueError(f"OCR 读取文件失败：{exc}") from exc
    payload["fileType"] = file_type
    payload["visualize"] = False
    payload["useDocUnwarping"] = False
    payload["useDocOrientationClassify"] = False

    headers = {
        "Authorization": f"token {access_token}",
        "Content-Type": "application/json",
        "Client-Platform": "bid-platform",
    }

    try:
        response = requests.post(api_url, json=payload, headers=headers, timeout=timeout_seconds)
    except requests.RequestException as exc:
        raise ValueError(f"OCR 服务连接失败：{exc}") from exc

    if response.status_code != 200:
        detail = response.text[:300]
        if response.status_code == 403:
            raise ValueError(f"OCR 鉴权失败（403）：{detail}")
        if response.status_code == 429:
            raise ValueError(f"OCR 调用频率或额度超限（429）：{detail}")
        raise ValueError(f"OCR 服务请求失败（HTTP {response.status_code}）：{detail}")

    try:
        response_payload = response.json()
    except ValueError as exc:
        raise ValueError(f"OCR 服务返回了非 JSON 内容：{response.text[:300]}") from exc

    if not isinstance(response_payload, dict):
        raise ValueError("OCR 服务返回格式异常。")
    if response_payload.get("errorCode", 0) != 0:
        raise ValueError(f"OCR 服务返回错误：{response_payload.get('errorMsg') or '未知错误'}")

    raw_result = response_payload.get("result")
    pages = raw_result.get("layoutParsingResults") if isinstance(raw_result, dict) else None
    if not isinstance(pages, list):
        raise ValueError("OCR 服务未返回可用版面结果。")

    parts: list[str] = []
    for index, page in enumerate(pages, start=1):
        markdown = page.get("markdown") if isinstance(page, dict) else None
        text = markdown.get("text") if isinstance(markdown, dict) else None
        if isinstance(text, str) and text.strip():
            parts.append(f"[第 {index} 页 OCR]\n{text.strip()}")

    if not parts:
        raise ValueError("OCR 完成，但未提取到可用文本。")
    return "\n\n".join(parts)


def extract_pdf_text(path: Path, ocr_settings: dict[str, object] | None = None) -> str:
    pages: list[str] = []
    try:
        with pdfplumber.open(path) as pdf:
            for index, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(f"[第 {index} 页]\n{text.strip()}")
        if pages:
            return "\n\n".join(pages)
    except Exception:
        pass

    try:
        reader = PdfReader(str(path))
        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"[第 {index} 页]\n{text.strip()}")
    except Exception as exc:
        raise ValueError(f"PDF 文本提取失败：{exc}") from exc

    if not pages:
        return extract_ocr_text(path, path.name, ocr_settings)
    return "\n\n".join(pages)


def summarize_docx_paragraph_style(paragraph) -> str:
    style_name = ""
    try:
        style_name = str(paragraph.style.name or "").strip()
    except Exception:
        style_name = ""
    if not style_name:
        return ""

    style_lower = style_name.lower()
    if "heading" in style_lower:
        match = re.search(r"(\d+)", style_name)
        if match:
            return f"标题 {match.group(1)}"
        return "标题"
    if "title" in style_lower:
        return "标题"
    if "subtitle" in style_lower:
        return "副标题"
    return ""


def append_docx_table_text(parts: list[str], table: Table, table_index: int, prefix: str = "表格") -> None:
    parts.append(f"[{prefix} {table_index}]")
    for row_index, row in enumerate(table.rows, start=1):
        cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells]
        if any(cells):
            parts.append(f"[{prefix} {table_index} | 第 {row_index} 行] " + " | ".join(cells))


def append_docx_header_footer_text(parts: list[str], document: Any) -> None:
    seen: set[str] = set()
    table_index = 0
    for section_index, section in enumerate(document.sections, start=1):
        for part_name, part in (
            ("页眉", section.header),
            ("首页页眉", section.first_page_header),
            ("偶数页页眉", section.even_page_header),
            ("页脚", section.footer),
            ("首页页脚", section.first_page_footer),
            ("偶数页页脚", section.even_page_footer),
        ):
            paragraph_index = 0
            for paragraph in part.paragraphs:
                text = paragraph.text.strip()
                if not text or text in seen:
                    continue
                seen.add(text)
                paragraph_index += 1
                parts.append(f"[第 {section_index} 节 | {part_name} | 段落 {paragraph_index}] {text}")
            for table in part.tables:
                table_index += 1
                append_docx_table_text(parts, table, table_index, prefix=f"第 {section_index} 节 {part_name}表格")


def flatten_docx2python_lines(value: object) -> list[str]:
    lines: list[str] = []
    if isinstance(value, str):
        text = value.strip()
        if text:
            lines.append(text)
        return lines
    if isinstance(value, (list, tuple)):
        for item in value:
            lines.extend(flatten_docx2python_lines(item))
    return lines


def extract_docx2python_supplement(path: Path, existing_text: str) -> list[str]:
    try:
        from docx2python import docx2python
    except ImportError:
        return []

    existing_lines = {re.sub(r"\s+", "", line) for line in existing_text.splitlines() if line.strip()}
    parts: list[str] = []
    try:
        content = docx2python(path)
    except Exception:
        return []

    try:
        supplement_sources = (
            ("页眉补充", getattr(content, "header", None)),
            ("页脚补充", getattr(content, "footer", None)),
            ("脚注补充", getattr(content, "footnotes", None)),
            ("尾注补充", getattr(content, "endnotes", None)),
            ("批注补充", getattr(content, "comments", None)),
            ("正文补充", getattr(content, "body", None)),
        )
        supplement_index = 0
        for source_name, source_value in supplement_sources:
            for line in flatten_docx2python_lines(source_value):
                normalized = re.sub(r"\s+", "", line)
                if not normalized or normalized in existing_lines:
                    continue
                existing_lines.add(normalized)
                supplement_index += 1
                parts.append(f"[docx2python | {source_name} | {supplement_index}] {line}")
                if supplement_index >= 500:
                    return parts
    finally:
        close = getattr(content, "close", None)
        if callable(close):
            close()

    return parts


def extract_docx_text(path: Path) -> str:
    if path.suffix.lower() == ".doc":
        raise ValueError("暂不支持旧版 .doc 自动解析，请先另存为 .docx 或 PDF 后上传。")

    document = Document(str(path))
    parts: list[str] = []
    paragraph_index = 0
    table_index = 0
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, document)
            text = paragraph.text.strip()
            if not text:
                continue
            paragraph_index += 1
            markers = [f"段落 {paragraph_index}"]
            style_summary = summarize_docx_paragraph_style(paragraph)
            if style_summary:
                markers.append(style_summary)
            parts.append(f"[{' | '.join(markers)}] {text}")
        elif isinstance(child, CT_Tbl):
            table_index += 1
            append_docx_table_text(parts, Table(child, document), table_index)

    append_docx_header_footer_text(parts, document)
    supplement_parts = extract_docx2python_supplement(path, "\n".join(parts))
    if supplement_parts:
        parts.append("[补充解析 | docx2python]")
        parts.extend(supplement_parts)

    return "\n".join(parts)


def extract_plain_text_with_line_markers(file_path: Path) -> str:
    content = file_path.read_text(encoding="utf-8", errors="ignore")
    lines = []
    for index, line in enumerate(content.splitlines(), start=1):
        if line.strip():
            lines.append(f"[第 {index} 行] {line.strip()}")
    return "\n".join(lines)


def extract_document_text(
    file_path: Path,
    original_name: str,
    ocr_settings: dict[str, object] | None = None,
) -> str:
    suffix = Path(original_name).suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_text(file_path, ocr_settings)
    if suffix in {".docx", ".doc"}:
        return extract_docx_text(file_path)
    if suffix in {".txt", ".md", ".csv"}:
        return extract_plain_text_with_line_markers(file_path)
    raise ValueError("当前仅支持 .docx、.doc、.pdf、.txt、.md、.csv 文件。")


def find_source_snippet(text: str, keywords: list[str], *, max_chars: int = 300) -> tuple[str, str]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for index, line in enumerate(lines):
        normalized_line = line.replace("\ufeff", "").strip()
        if any(keyword in normalized_line for keyword in keywords):
            snippet = "\n".join(lines[index : index + 3])[:max_chars]
            location_match = re.match(r"^\[(.+?)\]", normalized_line)
            if location_match:
                return location_match.group(1), snippet
            return f"第 {index + 1} 行附近", snippet
    return "", ""


def fallback_extract_project_payload(text: str, original_name: str) -> dict[str, object]:
    compact = re.sub(r"\s+", " ", text)
    name = Path(original_name).stem

    name_patterns = [
        r"(?:项目名称|采购项目名称|招标项目名称)[:：]\s*([^\r\n，。；;]{4,80})",
        r"([^\s，。；;]{4,80}(?:物业|服务|选聘|招标|采购)[^\s，。；;]{0,40})",
    ]
    for pattern in name_patterns:
        source = text if "项目名称" in pattern else compact
        match = re.search(pattern, source)
        if match:
            name = match.group(1).strip()
            break

    tender_code = None
    code_match = re.search(r"(?:项目编号|招标编号|采购编号)[:：]?\s*([A-Za-z0-9\-[\]【】（）()_.]+)", compact)
    if code_match:
        tender_code = code_match.group(1).strip()

    amount = None
    amount_match = re.search(r"保证金[^0-9一二三四五六七八九十百千万]{0,12}(\d+(?:\.\d+)?)\s*万", compact)
    if amount_match:
        amount = amount_match.group(1)

    _, invalidation_risks = find_source_snippet(text, FIELD_SOURCE_KEYWORDS["invalidation_risks"], max_chars=500)
    _, submission_notes = find_source_snippet(text, FIELD_SOURCE_KEYWORDS["submission_notes"], max_chars=500)
    _, seal_notes = find_source_snippet(text, FIELD_SOURCE_KEYWORDS["seal_notes"], max_chars=500)
    _, defense_snippet = find_source_snippet(text, FIELD_SOURCE_KEYWORDS["has_defense"], max_chars=300)
    has_defense = ""
    if defense_snippet:
        has_defense = "no" if re.search(r"无需|不需要|不组织|无答辩|无需答辩", defense_snippet) else "yes"

    return normalize_ai_project_payload(
        {
            "name": name,
            "short_name": name[:18],
            "bid_mode": "self",
            "deposit_amount": amount,
            "tender_code": tender_code,
            "has_defense": has_defense,
            "invalidation_risks": invalidation_risks,
            "submission_notes": submission_notes,
            "seal_notes": seal_notes,
            "notes": "系统未调用 AI，仅按文件名和少量规则生成草稿；请人工核对后保存。",
        }
    )


def build_fallback_field_sources(text: str, payload: dict[str, object]) -> dict[str, dict[str, str]]:
    sources: dict[str, dict[str, str]] = {}
    for field_key, keywords in FIELD_SOURCE_KEYWORDS.items():
        value = clean_text(payload.get(field_key))
        if not value:
            continue
        source_location, source_quote = find_source_snippet(text, keywords, max_chars=500)
        if not source_location and not source_quote:
            continue
        sources[field_key] = {
            "source_location": source_location,
            "source_quote": source_quote,
        }
    return sources


def fallback_extract_requirements(text: str) -> list[dict[str, str]]:
    rules: list[tuple[str, str, list[str]]] = [
        ("qualification", "资格要求", ["资格要求", "投标人资格", "申请人资格", "供应商资格"]),
        ("document", "投标文件组成", ["投标文件组成", "响应文件组成", "投标文件包括", "资格证明材料"]),
        ("technical", "技术要求", ["技术要求", "服务要求", "服务方案", "技术标准"]),
        ("business", "商务要求", ["商务要求", "商务条款", "合同条款", "付款方式"]),
        ("pricing", "报价要求", ["报价要求", "报价方式", "报价文件", "最高限价"]),
        ("scoring", "评分办法", ["评分办法", "评审办法", "评标办法", "综合评分"]),
        ("risk", "废标风险", ["废标", "无效投标", "否决投标", "无效响应"]),
    ]
    compact_lines = [line.strip() for line in text.splitlines() if line.strip()]
    results: list[dict[str, str]] = []
    for category, title, keywords in rules:
        for index, line in enumerate(compact_lines):
            if not any(keyword in line for keyword in keywords):
                continue
            snippet = "\n".join(compact_lines[index : index + 4])
            results.append(
                {
                    "category": category,
                    "category_label": REQUIREMENT_CATEGORY_LABELS.get(category, category),
                    "title": title,
                    "content": snippet[:300],
                    "importance": "medium",
                    "source_location": re.match(r"^\[(.+?)\]", line).group(1) if re.match(r"^\[(.+?)\]", line) else f"第 {index + 1} 行附近",
                    "source_quote": snippet[:300],
                }
            )
            break
    return results


def deepseek_chat(
    settings: dict[str, object],
    messages: list[dict[str, str]],
    *,
    json_output: bool = False,
    max_tokens: int = 4096,
) -> str:
    api_key = clean_text(settings.get("api_key"))
    if not api_key:
        raise ValueError("请先配置 DeepSeek API Key。")

    base_url = clean_text(settings.get("base_url")) or get_deepseek_base_url()
    model = clean_text(settings.get("model")) or get_deepseek_model()
    payload: dict[str, object] = {
        "model": model,
        "messages": messages,
        "temperature": 0.1,
        "max_tokens": max_tokens,
    }
    if json_output:
        payload["response_format"] = {"type": "json_object"}

    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    request = urllib.request.Request(
        f"{base_url.rstrip('/')}/chat/completions",
        data=body,
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )

    try:
        with urllib.request.urlopen(request, timeout=90) as response:
            response_payload = json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="ignore")
        raise ValueError(f"DeepSeek 请求失败：HTTP {exc.code} {detail[:500]}") from exc
    except urllib.error.URLError as exc:
        raise ValueError(f"DeepSeek 连接失败：{exc.reason}") from exc

    try:
        return response_payload["choices"][0]["message"]["content"]
    except (KeyError, IndexError, TypeError) as exc:
        raise ValueError("DeepSeek 返回格式异常。") from exc


def ping_deepseek(api_key: str | None = None) -> dict[str, str]:
    settings = get_runtime_ai_settings(api_key)
    if not settings["configured"]:
        raise ValueError("请先配置 DeepSeek API Key。")

    content = deepseek_chat(
        settings,
        [
            {"role": "system", "content": "你是连通性检测助手。"},
            {"role": "user", "content": "请只回复：ok"},
        ],
        json_output=False,
        max_tokens=20,
    ).strip()
    return {
        "status": "ok",
        "model": str(settings.get("model") or ""),
        "base_url": str(settings.get("base_url") or ""),
        "response_preview": content[:80],
    }


def extract_json_object(text: str) -> dict[str, object]:
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?", "", cleaned, flags=re.IGNORECASE).strip()
        cleaned = re.sub(r"```$", "", cleaned).strip()
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", cleaned, flags=re.DOTALL)
        if not match:
            raise ValueError("AI 返回内容不是合法 JSON。")
        return json.loads(match.group(0))


def collect_keyword_windows(
    text: str,
    keywords: list[str],
    *,
    before: int = 2,
    after: int = 6,
    max_windows: int = 180,
) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    ranges: list[tuple[int, int]] = []
    for index, line in enumerate(lines):
        if not any(keyword in line for keyword in keywords):
            continue
        start = max(0, index - before)
        end = min(len(lines), index + after + 1)
        if ranges and start <= ranges[-1][1]:
            ranges[-1] = (ranges[-1][0], max(ranges[-1][1], end))
        else:
            ranges.append((start, end))
        if len(ranges) >= max_windows:
            break

    parts: list[str] = []
    for window_index, (start, end) in enumerate(ranges, start=1):
        parts.append(f"[重点窗口 {window_index}]")
        parts.extend(lines[start:end])
    return "\n".join(parts)


def build_model_extraction_context(extracted_text: str) -> str:
    if len(extracted_text) <= MAX_MODEL_CONTEXT_CHARS:
        return extracted_text

    head = extracted_text[:MODEL_CONTEXT_HEAD_CHARS]
    tail = extracted_text[-MODEL_CONTEXT_TAIL_CHARS:]
    remaining = max(MAX_MODEL_CONTEXT_CHARS - len(head) - len(tail) - 2000, 0)
    keyword_windows = collect_keyword_windows(extracted_text, DOCUMENT_CONTEXT_KEYWORDS)
    keyword_windows = keyword_windows[:remaining]
    return (
        f"{head}\n\n"
        "--- 全文重点条款候选（按字段关键词从全文抓取，避免长文档截断遗漏）---\n"
        f"{keyword_windows}\n\n"
        "--- 文末候选（常见附件、格式、废标条款位置）---\n"
        f"{tail}"
    )[:MAX_MODEL_CONTEXT_CHARS]


def enrich_requirement_sources_from_text(items: list[dict[str, str]], extracted_text: str) -> list[dict[str, str]]:
    enriched: list[dict[str, str]] = []
    for item in items:
        normalized = dict(item)
        source_quote = clean_text(normalized.get("source_quote")) or ""
        if source_quote and quote_is_grounded_in_text(source_quote, extracted_text):
            enriched.append(normalized)
            continue

        category = clean_text(normalized.get("category")) or "other"
        title = clean_text(normalized.get("title")) or ""
        keywords = [title, *REQUIREMENT_SOURCE_KEYWORDS.get(category, []), *FIELD_SOURCE_KEYWORDS.get("invalidation_risks", [])]
        keywords = [keyword for keyword in keywords if keyword]
        source_location, fallback_quote = find_source_snippet(extracted_text, keywords, max_chars=500)
        if fallback_quote:
            normalized["source_location"] = source_location
            normalized["source_quote"] = fallback_quote
        elif source_quote and not quote_is_grounded_in_text(source_quote, extracted_text):
            normalized["source_location"] = ""
            normalized["source_quote"] = ""
        enriched.append(normalized)
    return enriched


def ai_extract_fields(
    file_name: str,
    extracted_text: str,
    api_key: str | None = None,
) -> tuple[dict[str, object], list[dict[str, str]], list[dict[str, str]], str]:
    def build_rule_fallback_result(reason: str) -> tuple[dict[str, object], list[dict[str, str]], list[dict[str, str]], str]:
        fallback = fallback_extract_project_payload(extracted_text, file_name)
        fallback_sources = build_fallback_field_sources(extracted_text, fallback)
        fallback_fields = []
        for field_key, field_label in FIELD_LABELS.items():
            value = fallback.get(field_key)
            source_meta = fallback_sources.get(field_key, {})
            fallback_fields.append(
                {
                    "field_key": field_key,
                    "field_label": field_label,
                    "extracted_value": "" if value is None else str(value),
                    "source_location": source_meta.get("source_location", ""),
                    "source_quote": source_meta.get("source_quote", ""),
                    "confidence": "medium" if field_key in fallback_sources else "low",
                }
            )
        fallback_requirements = fallback_extract_requirements(extracted_text)
        summary = f"{reason}；系统已生成规则兜底解析，请人工核对后再入库。"
        return fallback, fallback_fields, fallback_requirements, summary

    settings = get_runtime_ai_settings(api_key)
    if not settings["configured"]:
        return build_rule_fallback_result("当前未配置 DeepSeek")

    model_context = build_model_extraction_context(extracted_text)
    prompt = f"""
你是投标项目解析助手。请基于招标文件原文提取项目关键信息，并返回一个 JSON 对象。

JSON 结构要求：
{{
  "summary": "一句话总结解析情况",
  "project": {{
    "name": "",
    "short_name": "",
    "tender_code": "",
    "buyer": "",
    "project_type": "",
    "bid_mode": "self 或 partner",
    "owner_name": "待分配 或原文信息",
    "agency": "",
    "contact_name": "",
    "contact_phone": "",
    "location": "",
    "service_scope": "",
    "contract_term": "",
    "budget_amount": "",
    "deposit_amount": "",
    "signup_deadline": "YYYY-MM-DD HH:MM 或空字符串",
    "document_sale_deadline": "YYYY-MM-DD HH:MM 或空字符串",
    "clarification_deadline": "YYYY-MM-DD HH:MM 或空字符串",
    "site_visit_time": "YYYY-MM-DD HH:MM 或空字符串",
    "deposit_deadline": "YYYY-MM-DD HH:MM 或空字符串",
    "bid_datetime": "YYYY-MM-DD HH:MM 或空字符串",
    "submission_datetime": "YYYY-MM-DD HH:MM 或空字符串",
    "bid_location": "",
    "file_fee": "",
    "payment_info": "",
    "has_defense": "yes / no / 空字符串；只有原文明确提到答辩、述标、演示时才填写",
    "defense_presenter": "例如 项目负责人/项目经理/授权代表；原文未写则空",
    "invalidation_risks": "废标项、无效投标、否决投标风险摘要",
    "submission_notes": "递交注意事项，例如份数、地点、携带材料、送达要求",
    "seal_notes": "封标注意事项，例如密封方式、封套/封条/封皮内容、盖章要求",
    "notes": ""
  }},
  "fields": [
    {{
      "field_key": "name",
      "field_label": "项目名称",
      "extracted_value": "值",
      "source_location": "例如 第 3 页 / 表格 2 / 段落 5",
      "source_quote": "原文摘录",
      "confidence": "high / medium / low"
    }}
  ],
  "requirements": [
    {{
      "category": "qualification / document / business / technical / pricing / scoring / risk / response / other",
      "title": "简短标题",
      "content": "要求内容摘要",
      "importance": "high / medium / low",
      "source_location": "例如 第 8 页 / 表格 3 / 段落 12",
      "source_quote": "原文摘录"
    }}
  ]
}}

规则：
1. 只能依据原文，不允许编造。
2. 找不到的字段填空字符串，并把 confidence 设为 low。
3. 每个有值字段都必须附 source_location 和 source_quote；source_quote 必须是原文逐字摘录，不要改写，单条不超过 220 字。
4. requirements 至少尽量覆盖：资格要求、投标文件组成、评分办法、报价要求、废标风险、关键响应点。
5. 特别提取并归纳：递交注意事项、封标注意事项、废标项/无效投标风险、是否有答辩及答辩人要求。
6. 如果某个字段写进 fields，也要同步写进 project 对应字段。
7. requirements 最多返回 24 条；每个 category 选最关键的 1-4 条即可，不要整段复制长篇原文。
8. 如果某一类没找到，不必强行编造。
9. 只返回合法 JSON，不要写 Markdown。

文件名：{file_name}
招标文件原文：
{model_context}
"""
    try:
        content = deepseek_chat(
            settings,
            [
                {"role": "system", "content": "你是一名严谨的投标文件解析助手，只能输出合法 JSON。"},
                {"role": "user", "content": prompt},
            ],
            json_output=True,
            max_tokens=8000,
        )
    except ValueError as exc:
        return build_rule_fallback_result(f"DeepSeek 调用失败：{str(exc)[:300]}")

    try:
        payload = extract_json_object(content)
    except (json.JSONDecodeError, ValueError) as exc:
        return build_rule_fallback_result(f"DeepSeek 返回 JSON 格式异常：{str(exc)[:300]}")
    raw_fields = payload.get("fields", [])
    project_payload = merge_project_payload_from_fields(payload.get("project", {}), raw_fields)
    normalized_project = normalize_ai_project_payload(project_payload)
    fields = normalize_ai_field_payload(raw_fields, normalized_project, extracted_text)
    requirements = enrich_requirement_sources_from_text(normalize_requirement_payload(payload.get("requirements", [])), extracted_text)
    summary = str(payload.get("summary", "")).strip() or "已完成 AI 提取，请人工核对。"
    return normalized_project, fields, requirements, summary


def build_project_context(
    project_name: str,
    project_status_label: str,
    files: list[object],
    requirements: list[object] | None = None,
) -> str:
    context_parts = [f"项目名称：{project_name}", f"状态：{project_status_label}"]
    if requirements:
        context_parts.append("\n--- 结构化投标准备要点 ---")
        for item in requirements:
            category = getattr(item, "category", "other")
            title = getattr(item, "title", "")
            content = getattr(item, "content", "")
            source_location = getattr(item, "source_location", "")
            context_parts.append(
                f"[{REQUIREMENT_CATEGORY_LABELS.get(category, category)}] {title}\n{content}\n来源：{source_location or '未记录'}"
            )
    for file_record in files:
        if not getattr(file_record, "extracted_text", None):
            continue
        context_parts.append(f"\n--- 文件：{file_record.original_name} ---\n{file_record.extracted_text}")
    return "\n".join(context_parts)[:MAX_MODEL_CONTEXT_CHARS]


def extract_value_from_quote(source_quote: str) -> str:
    quote = clean_text(source_quote) or ""
    if "：" in quote:
        return quote.split("：", 1)[1].strip() or quote
    if ":" in quote:
        return quote.split(":", 1)[1].strip() or quote
    return quote


def collect_context_matches(project_context: str, hints: list[str], *, limit: int = 2) -> list[dict[str, str]]:
    lines = [line.strip() for line in project_context.splitlines() if line.strip()]
    current_file = ""
    matches: list[dict[str, str]] = []
    seen_keys: set[tuple[str, str, str]] = set()

    for index, line in enumerate(lines):
        if line.startswith("--- 文件：") and line.endswith(" ---"):
            current_file = line.removeprefix("--- 文件：").removesuffix(" ---").strip()
            continue

        window_lines = lines[index : index + 3]
        matching_offset = next(
            (
                offset
                for offset, candidate_line in enumerate(window_lines)
                if any(hint in candidate_line for hint in hints)
            ),
            None,
        )
        if matching_offset is None:
            continue
        matched_index = index + matching_offset
        matched_line = lines[matched_index]

        source_location = ""
        source_quote = ""
        explanation = ""
        file_name = current_file

        if matched_line.startswith("[") and "]" in matched_line:
            source_location = matched_line[1:].split("]", 1)[0].strip()
            source_quote = matched_line.split("]", 1)[1].strip()
            explanation = "以下原文直接支持本次回答。"
        else:
            source_line = matched_line
            if matched_index + 1 < len(lines) and lines[matched_index + 1] and not lines[matched_index + 1].startswith("来源："):
                source_line = lines[matched_index + 1]
            source_quote = source_line.strip()
            if matched_index + 1 < len(lines) and lines[matched_index + 1].startswith("来源："):
                source_location = lines[matched_index + 1].split("：", 1)[1].strip()
            elif matched_index + 2 < len(lines) and lines[matched_index + 2].startswith("来源："):
                source_location = lines[matched_index + 2].split("：", 1)[1].strip()
            explanation = "结构化要点或文件原文与问题直接相关。"

        if not source_quote:
            continue

        key = (file_name, source_location, source_quote)
        if key in seen_keys:
            continue
        seen_keys.add(key)
        matches.append(
            {
                "file_name": file_name,
                "source_location": source_location,
                "source_quote": source_quote,
                "explanation": explanation,
            }
        )
        if len(matches) >= limit:
            break

    return matches


def answer_project_question_locally(question: str, project_context: str) -> tuple[str, str]:
    normalized_question = repair_mojibake_text(clean_text(question) or "")
    matched_groups = [group for group in LOCAL_QA_GROUPS if any(hint in normalized_question for hint in group["hints"])]

    answer_parts: list[str] = []
    basis_items: list[dict[str, str]] = []
    missing_labels: list[str] = []
    seen_basis: set[tuple[str, str, str]] = set()

    for group in matched_groups:
        matches = collect_context_matches(project_context, group["hints"], limit=1 if group["mode"] == "value" else 2)
        if not matches:
            missing_labels.append(group["label"])
            continue

        primary = matches[0]
        if group["mode"] == "value":
            answer_parts.append(f"{group['label']}：{extract_value_from_quote(primary['source_quote'])}")
        else:
            answer_parts.append(f"{group['label']}：{primary['source_quote']}")

        for item in matches:
            basis_key = (item["file_name"], item["source_location"], item["source_quote"])
            if basis_key in seen_basis:
                continue
            seen_basis.add(basis_key)
            basis_items.append(item)
            if len(basis_items) >= 4:
                break
        if len(basis_items) >= 4:
            break

    if answer_parts:
        answer_status = "grounded" if not missing_labels else "partial"
        answer = "；".join(answer_parts)
        if missing_labels:
            answer += f"；其余关于{'、'.join(missing_labels)}的信息，当前文件依据不足。"
        payload = normalize_question_answer_payload(
            {
                "answer_mode": "fallback",
                "answer_status": answer_status,
                "answer": answer,
                "basis_items": basis_items,
                "ai_suggestion": "" if answer_status == "grounded" else "如需更复杂的综合判断，建议后续配置 DeepSeek API Key。",
            }
        )
        return str(payload["answer"]), serialize_question_answer_metadata(payload)

    not_found_payload = normalize_question_answer_payload(
        {
            "answer_mode": "fallback",
            "answer_status": "not_found",
            "answer": "当前已上传的文件中未找到可以直接支持该问题的明确依据。",
            "basis_items": [],
            "ai_suggestion": "可以先补充相关附件、答疑文件，或配置 DeepSeek API Key 后再提问。",
        }
    )
    return str(not_found_payload["answer"]), serialize_question_answer_metadata(not_found_payload)


def answer_project_question(question: str, project_context: str, api_key: str | None = None) -> tuple[str, str]:
    question = repair_mojibake_text(clean_text(question) or "")
    settings = get_runtime_ai_settings(api_key)
    if not settings["configured"]:
        return answer_project_question_locally(question, project_context)

    prompt = f"""
你是投标项目问答助手。你只能依据提供的项目资料回答，绝对不能编造。

请返回一个 JSON 对象，结构如下：
{{
  "answer_status": "grounded / partial / not_found",
  "answer": "对用户问题的直接回答，必须简洁明确",
  "basis_items": [
    {{
      "file_name": "来源文件名",
      "source_location": "例如 第 8 页 / 表格 3 / 结构化要点",
      "source_quote": "支持回答的原文摘录",
      "explanation": "这段原文如何支持答案"
    }}
  ],
  "ai_suggestion": "仅当原文未明确写出、但可给用户下一步建议时填写；如果没有可留空"
}}

规则：
1. answer 只能写文件里能支持的内容。
2. 如果资料只能支持部分答案，answer_status 填 partial，并明确哪些已知、哪些未知。
3. 如果资料里没有明确答案，answer_status 填 not_found，answer 要明确说明未找到。
4. basis_items 最多 4 条，每条尽量附原文摘录。
5. ai_suggestion 只能是流程建议或补充资料建议，不能冒充文件事实。
6. 只返回合法 JSON，不要输出 Markdown。

项目资料：
{project_context}

用户问题：
{question}
"""

    content = deepseek_chat(
        settings,
        [
            {
                "role": "system",
                "content": "你是一名严谨的投标项目问答助手，只能输出合法 JSON，不能编造文件中没有的内容。",
            },
            {"role": "user", "content": prompt},
        ],
        json_output=True,
        max_tokens=4000,
    )
    payload = extract_json_object(content)
    payload["answer_mode"] = "deepseek"
    normalized = normalize_question_answer_payload(payload)
    return str(normalized["answer"]), serialize_question_answer_metadata(normalized)


def format_datetime(value: datetime | None) -> str:
    if value is None:
        return ""
    return value.strftime("%Y-%m-%d %H:%M")
